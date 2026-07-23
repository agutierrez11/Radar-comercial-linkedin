import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml

def create_strict_cv(output_path, lang='es'):
    doc = docx.Document()
    
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        
    PRIMARY_COLOR = RGBColor(15, 23, 42)      # #0f172a
    ACCENT_COLOR = RGBColor(37, 99, 235)      # #2563eb
    TEXT_COLOR = RGBColor(51, 65, 85)         # #334155
    MUTED_COLOR = RGBColor(100, 116, 139)     # #64748b
    
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = TEXT_COLOR
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(4)

    def add_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = ACCENT_COLOR
        
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="12" w:space="4" w:color="2563EB"/></w:pBdr>')
        pPr.append(pBdr)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.name = 'Calibri'
            run_b.font.size = Pt(10)
            run_b.font.bold = True
            run_b.font.color.rgb = PRIMARY_COLOR
            
        run_t = p.add_run(text)
        run_t.font.name = 'Calibri'
        run_t.font.size = Pt(10)
        run_t.font.color.rgb = TEXT_COLOR
        return p

    if lang == 'es':
        # HEADER ES
        p_name = doc.add_paragraph()
        p_name.paragraph_format.space_after = Pt(2)
        r_name = p_name.add_run("ANTONIO GUTIÉRREZ JIMÉNEZ")
        r_name.font.name = 'Calibri'
        r_name.font.size = Pt(20)
        r_name.font.bold = True
        r_name.font.color.rgb = PRIMARY_COLOR

        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_after = Pt(6)
        r_title = p_title.add_run("B2B Sales Executive  |  Fintech & Payments Senior Specialist")
        r_title.font.name = 'Calibri'
        r_title.font.size = Pt(11)
        r_title.font.bold = True
        r_title.font.color.rgb = ACCENT_COLOR

        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(10)
        r_contact = p_contact.add_run("📍 Cancún, Quintana Roo, México   |   📞 +52 998 119 1903   |   ✉️ antoniogtzjimenez@gmail.com   |   🔗 linkedin.com/in/agjbusiness/")
        r_contact.font.name = 'Calibri'
        r_contact.font.size = Pt(9.5)
        r_contact.font.color.rgb = MUTED_COLOR

        # RESUMEN PROFESIONAL
        add_heading("Resumen Profesional")
        p_sum = doc.add_paragraph()
        p_sum.add_run(
            "Ejecutivo de Ventas B2B Senior y especialista en Adquirencia (Merchant Acquiring) y Tecnología de Pagos (PayTech), con más de 5 años de trayectoria acelerando el crecimiento comercial de grandes corporaciones fintech y multinacionales en México y LATAM. Reconocido por un enfoque consultivo de alta eficiencia: priorizo la adquisición de cuentas estratégicas (Middle Market/High Potential) de alto volumen transaccional, optimizando el costo de onboarding e integraciones tecnológicas (APIs/ISVs). Experto en prospección outbound autónoma en campo y corporativa, con un portafolio activo de más de 3,000 conexiones en la industria de pagos digitales en LATAM."
        )

        # ÁREAS DE EXPERTISE
        add_heading("Áreas de Expertise")
        add_bullet("Procesamiento de pagos, adquirencia, pasarelas de pago (gateways), prevención de fraude, transacciones transfronterizas.", bold_prefix="Merchant Acquiring & PayTech: ")
        add_bullet("Prospección activa, mapeo de territorios, ciclos de venta complejos B2B de ciclo largo.", bold_prefix="Ventas Consultivas & Outbound Hunting: ")
        add_bullet("Alianzas comerciales y técnicas con ERPs, PMS y sistemas de punto de venta (POS) para automatización de comercios.", bold_prefix="Integraciones API & ISV Partnerships: ")
        add_bullet("Creación de herramientas propias de SalesTech, automatizaciones con Power BI / Power Automate, Salesforce y CRM.", bold_prefix="Sales Ops & Inteligencia Comercial: ")

        # EXPERIENCIA PROFESIONAL
        add_heading("Experiencia Profesional")

        # Co-Founder
        p_e1 = doc.add_paragraph()
        p_e1.paragraph_format.space_before = Pt(6)
        p_e1.paragraph_format.space_after = Pt(2)
        p_e1.paragraph_format.keep_with_next = True
        r_role = p_e1.add_run("Co-Founder  |  LATAM Payments & eCommerce")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e1.add_run("   •   05/2024 – Presente   •   LATAM · Remoto")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Cofundé y coordino una comunidad activa de más de 500 profesionales de pagos digitales, adquirencia y comercio electrónico en México y LATAM.")

        # Fiserv
        p_e2 = doc.add_paragraph()
        p_e2.paragraph_format.space_before = Pt(8)
        p_e2.paragraph_format.space_after = Pt(2)
        p_e2.paragraph_format.keep_with_next = True
        r_role = p_e2.add_run("Business Advisor  |  Fiserv")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e2.add_run("   •   02/2025 – 10/2025   •   Cancún, México")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Gestioné la retención, mitigación de Churn y desarrollo de una cartera activa de más de 80 comercios corporativos (merchants) de gran escala.")
        add_bullet("Diseñé e implementé un modelo dinámico de monitoreo transaccional Month-over-Month (MoM) mediante Power BI, optimizando la toma de decisiones comerciales.")
        add_bullet("Desarrollé un workflow automatizado de reactivación mediante Power Automate e integraciones de Salesforce, generando un promedio de +15 oportunidades de negocio mensuales desde la alianza bancaria asignada.")

        # Clip
        p_e3 = doc.add_paragraph()
        p_e3.paragraph_format.space_before = Pt(8)
        p_e3.paragraph_format.space_after = Pt(2)
        p_e3.paragraph_format.keep_with_next = True
        r_role = p_e3.add_run("Asesor Comercial – Middle Market / High Potential  |  Clip")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e3.add_run("   •   07/2021 – 02/2025   •   Cancún, México")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Posicionado en el Top 12% nacional (Lugar #22 de 184 ejecutivos de Middle Market) en H1 2022, superando las cuotas de volumen mensual asignadas por más del 280% de forma consistente ($2.8M a $5.8M MXN promedio frente a la meta de $1M) y logrando el 3er lugar general del podio nacional.", bold_prefix="Top Performer & Cumplimiento: ")
        add_bullet("Diseñé y ejecuté una estrategia comercial enfocada en cuentas medianas de alto potencial, logrando un TPV promedio por deal de $555k MXN (60% superior a la media del segmento), maximizando el volumen procesado con una fracción del costo operativo de integración y soporte.", bold_prefix="Eficiencia de Cartera (High Value): ")
        add_bullet("Cerré de manera autónoma las cuentas de mayor volumen de la cartera en el sector turismo de lujo, destacando un operador de yates de lujo ($14.5M MXN YTD TPV) y una firma de turismo de aventura ($20.0M MXN YTD TPV) mediante prospección activa en frío y cambaceo estratégico.", bold_prefix="Cierre de Cuentas Enterprise (Outbound): ")
        add_bullet("Lideré negociaciones comerciales complejas e integraciones de pasarela de pagos vía API/ISV con sistemas clave (Bistrosoft, Profitroom, Odoo ERP), incrementando la retención de clientes a largo plazo con una tasa de churn cercana a cero.", bold_prefix="Integraciones Tecnológicas & APIs: ")
        add_bullet("$69M MXN de TPV total acumulado, con un 75.3% del volumen total auto-generado vía Outbound.", bold_prefix="Resultado consolidado de la cartera: ")

        # JTI
        p_e4 = doc.add_paragraph()
        p_e4.paragraph_format.space_before = Pt(8)
        p_e4.paragraph_format.space_after = Pt(2)
        p_e4.paragraph_format.keep_with_next = True
        r_role = p_e4.add_run("Account Executive · Southeast & Bajío  |  Japan Tobacco International (JTI)")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e4.add_run("   •   07/2018 – 12/2020   •   Cancún & Aguascalientes, México")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Lideré la expansión territorial y negociación de Key Accounts (KAM) en el sector HORECA, logrando un crecimiento del +40% en Share of Opportunity en Cancún y Riviera Maya vs. el año previo.")
        add_bullet("Expandí la cartera de clientes activos en un +35%, asegurando contratos de distribución directa con más de 100 hoteles premium y cadenas internacionales de hospitalidad líderes en la región.")
        add_bullet("Coordiné y lideré un equipo comercial de 3 personas en campo (FSF) para la región del Bajío.")

        # EDUCACIÓN
        add_heading("Educación")
        p_ed = doc.add_paragraph()
        p_ed.paragraph_format.space_after = Pt(2)
        r_ed = p_ed.add_run("Licenciatura en Relaciones Comerciales")
        r_ed.bold = True
        r_ed.font.color.rgb = PRIMARY_COLOR
        p_ed2 = doc.add_paragraph()
        p_ed2.paragraph_format.space_after = Pt(4)
        r_ed2 = p_ed2.add_run("Instituto Politécnico Nacional (IPN) — México   |   2014 – 2018 (Titulado)")
        r_ed2.font.color.rgb = MUTED_COLOR
        r_ed2.font.size = Pt(9.5)

        # CERTIFICACIONES
        add_heading("Educación Continua & Certificaciones")
        add_bullet("McKinsey Forward Program (McKinsey.org · 120 horas) — Especialización en liderazgo adaptativo, resolución estructurada de problemas complejos y metodologías ágiles.")
        add_bullet("Growth 101 (Kurios · 30 horas) — Metodología de crecimiento acelerado y marcos de experimentación ágil de producto.")
        add_bullet("Mastering Ventas (Sales Professional · 70 horas) — Estructura de equipos comerciales, playbook de ventas B2B y automatización del stack SalesTech.")
        add_bullet("Curso SDR – Primera Reunión (LATAM SDR Leaders · 16 horas) — Estrategias avanzadas de prospección en frío y agendamiento de cuentas Enterprise.")

    else:
        # HEADER EN (Exact Word-for-Word Translation of Base CV)
        p_name = doc.add_paragraph()
        p_name.paragraph_format.space_after = Pt(2)
        r_name = p_name.add_run("ANTONIO GUTIÉRREZ JIMÉNEZ")
        r_name.font.name = 'Calibri'
        r_name.font.size = Pt(20)
        r_name.font.bold = True
        r_name.font.color.rgb = PRIMARY_COLOR

        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_after = Pt(6)
        r_title = p_title.add_run("B2B Sales Executive  |  Fintech & Payments Senior Specialist")
        r_title.font.name = 'Calibri'
        r_title.font.size = Pt(11)
        r_title.font.bold = True
        r_title.font.color.rgb = ACCENT_COLOR

        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(10)
        r_contact = p_contact.add_run("📍 Cancún, Quintana Roo, Mexico   |   📞 +52 998 119 1903   |   ✉️ antoniogtzjimenez@gmail.com   |   🔗 linkedin.com/in/agjbusiness/")
        r_contact.font.name = 'Calibri'
        r_contact.font.size = Pt(9.5)
        r_contact.font.color.rgb = MUTED_COLOR

        # PROFESSIONAL SUMMARY
        add_heading("Professional Summary")
        p_sum = doc.add_paragraph()
        p_sum.add_run(
            "Senior B2B Sales Executive and specialist in Merchant Acquiring and Payment Technology (PayTech), with 5+ years of track record accelerating commercial growth for major fintech corporations and multinationals in Mexico and LATAM. Recognized for a high-efficiency consultative approach: prioritizing the acquisition of strategic accounts (Middle Market/High Potential) with high transactional volume, optimizing onboarding costs and technological integrations (APIs/ISVs). Expert in autonomous field and corporate outbound prospecting, with an active portfolio of over 3,000 connections in the digital payments industry in LATAM."
        )

        # AREAS OF EXPERTISE
        add_heading("Areas of Expertise")
        add_bullet("Payment processing, merchant acquiring, payment gateways, fraud prevention, cross-border transactions.", bold_prefix="Merchant Acquiring & PayTech: ")
        add_bullet("Active prospecting, territory mapping, complex long-cycle B2B sales cycles.", bold_prefix="Consultative Sales & Outbound Hunting: ")
        add_bullet("Commercial and technical partnerships with ERPs, PMS, and Point of Sale (POS) systems for merchant automation.", bold_prefix="API & ISV Partnerships: ")
        add_bullet("Development of proprietary SalesTech tools, automation with Power BI / Power Automate, Salesforce, and CRM.", bold_prefix="Sales Ops & Business Intelligence: ")

        # PROFESSIONAL EXPERIENCE
        add_heading("Professional Experience")

        # Co-Founder
        p_e1 = doc.add_paragraph()
        p_e1.paragraph_format.space_before = Pt(6)
        p_e1.paragraph_format.space_after = Pt(2)
        p_e1.paragraph_format.keep_with_next = True
        r_role = p_e1.add_run("Co-Founder  |  LATAM Payments & eCommerce")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e1.add_run("   •   05/2024 – Present   •   LATAM · Remote")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Co-founded and coordinate an active community of over 500 digital payments, acquiring, and e-commerce professionals in Mexico and LATAM.")

        # Fiserv
        p_e2 = doc.add_paragraph()
        p_e2.paragraph_format.space_before = Pt(8)
        p_e2.paragraph_format.space_after = Pt(2)
        p_e2.paragraph_format.keep_with_next = True
        r_role = p_e2.add_run("Business Advisor  |  Fiserv")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e2.add_run("   •   02/2025 – 10/2025   •   Cancún, Mexico")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Managed retention, churn mitigation, and business development for an active portfolio of over 80 large-scale corporate merchants.")
        add_bullet("Designed and implemented a dynamic Month-over-Month (MoM) transactional monitoring model using Power BI, optimizing commercial decision-making.")
        add_bullet("Developed an automated reactivation workflow using Power Automate and Salesforce integrations, generating an average of +15 monthly business opportunities from assigned banking alliances.")

        # Clip
        p_e3 = doc.add_paragraph()
        p_e3.paragraph_format.space_before = Pt(8)
        p_e3.paragraph_format.space_after = Pt(2)
        p_e3.paragraph_format.keep_with_next = True
        r_role = p_e3.add_run("Commercial Advisor – Middle Market / High Potential  |  Clip")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e3.add_run("   •   07/2021 – 02/2025   •   Cancún, Mexico")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Ranked in the Top 12% nationally (#22 out of 184 Middle Market executives) in H1 2022, consistently exceeding assigned monthly volume quotas by over 280% ($2.8M to $5.8M MXN average vs. $1M goal), achieving 3rd place overall on the national performance podium.", bold_prefix="Top Performer & Compliance: ")
        add_bullet("Designed and executed a commercial strategy focused on high-potential mid-market accounts, achieving an average TPV per deal of $555k MXN (60% above segment average), maximizing volume processed with a fraction of integration and support operational costs.", bold_prefix="Portfolio Efficiency (High Value): ")
        add_bullet("Independently closed the highest-volume accounts in the regional portfolio in luxury tourism, highlighting a luxury yacht operator ($14.5M MXN YTD TPV) and an adventure tourism firm ($20.0M MXN YTD TPV) through active cold prospecting and strategic field sales.", bold_prefix="Enterprise Account Closing (Outbound): ")
        add_bullet("Led complex commercial negotiations and payment gateway integrations via API/ISV with key systems (Bistrosoft, Profitroom, Odoo ERP), increasing long-term customer retention with near-zero churn.", bold_prefix="Technology & API Integrations: ")
        add_bullet("$69M MXN total cumulative TPV, with 75.3% of total volume self-generated via Outbound.", bold_prefix="Consolidated Portfolio Result: ")

        # JTI
        p_e4 = doc.add_paragraph()
        p_e4.paragraph_format.space_before = Pt(8)
        p_e4.paragraph_format.space_after = Pt(2)
        p_e4.paragraph_format.keep_with_next = True
        r_role = p_e4.add_run("Account Executive · Southeast & Bajío  |  Japan Tobacco International (JTI)")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e4.add_run("   •   07/2018 – 12/2020   •   Cancún & Aguascalientes, Mexico")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Led territorial expansion and Key Account Management (KAM) in the HORECA sector, achieving +40% growth in Share of Opportunity in Cancún and Riviera Maya vs. previous year.")
        add_bullet("Expanded active customer portfolio by +35%, securing direct distribution contracts with over 100 premium hotels and leading international hospitality chains.")
        add_bullet("Coordinated and led a field sales team of 3 people (FSF) for the Bajío region.")

        # EDUCATION
        add_heading("Education")
        p_ed = doc.add_paragraph()
        p_ed.paragraph_format.space_after = Pt(2)
        r_ed = p_ed.add_run("Bachelor’s Degree in Commercial Relations (Licenciatura en Relaciones Comerciales)")
        r_ed.bold = True
        r_ed.font.color.rgb = PRIMARY_COLOR
        p_ed2 = doc.add_paragraph()
        p_ed2.paragraph_format.space_after = Pt(4)
        r_ed2 = p_ed2.add_run("Instituto Politécnico Nacional (IPN) — Mexico   |   2014 – 2018 (Graduated / Titulado)")
        r_ed2.font.color.rgb = MUTED_COLOR
        r_ed2.font.size = Pt(9.5)

        # CERTIFICATIONS
        add_heading("Continuous Education & Certifications")
        add_bullet("McKinsey Forward Program (McKinsey.org · 120 hours) — Specialization in adaptive leadership, structured complex problem solving, and agile methodologies.")
        add_bullet("Growth 101 (Kurios · 30 hours) — Accelerated growth methodology and rapid product experimentation frameworks.")
        add_bullet("Mastering Ventas (Sales Professional · 70 hours) — B2B sales team structure, sales playbooks, and SalesTech stack automation.")
        add_bullet("SDR Course – First Meeting (LATAM SDR Leaders · 16 hours) — Advanced cold prospecting strategies and Enterprise account scheduling.")

    doc.save(output_path)
    print(f"Generated strict Word CV: {output_path}")

out_dir = r"c:\Users\Antonio\OneDrive\Escritorio\CVs_2026"
create_strict_cv(os.path.join(out_dir, "CV_Antonio_Gutierrez_MACHOTE.docx"), lang='es')
create_strict_cv(os.path.join(out_dir, "CV_Antonio_Gutierrez_MACHOTE_EN.docx"), lang='en')
