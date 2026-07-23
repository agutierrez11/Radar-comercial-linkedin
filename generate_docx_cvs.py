import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_styled_cv(output_path, lang='en'):
    doc = docx.Document()
    
    # Page setup - Margins 0.6 inch
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        
    # Color palette
    PRIMARY_COLOR = RGBColor(15, 23, 42)      # Dark Navy #0f172a
    ACCENT_COLOR = RGBColor(37, 99, 235)      # Royal Blue #2563eb
    TEXT_COLOR = RGBColor(51, 65, 85)         # Slate Gray #334155
    MUTED_COLOR = RGBColor(100, 116, 139)     # Light Slate #64748b
    
    # Base Normal Style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = TEXT_COLOR
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(4)

    # Helper function for Section Headings
    def add_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = ACCENT_COLOR
        
        # Add bottom border XML
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="12" w:space="4" w:color="2563EB"/></w:pBdr>')
        pPr.append(pBdr)
        return p

    def add_bullet(p_or_text, text=None, bold_prefix=None):
        if text is None:
            text = p_or_text
            p = doc.add_paragraph(style='List Bullet')
        else:
            p = p_or_text
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.name = 'Calibri'
            run_b.font.size = Pt(10)
            run_b.font.bold = True
            run_b.font.color.rgb = PRIMARY_COLOR
            
        run_t = p.add_run(text)
        run_t.font.name = 'Calibri'
        run_t.font.size = Pt(10)
        run_t.font.color.rgb = TEXT_COLOR
        return p

    if lang == 'en':
        # HEADER EN
        p_name = doc.add_paragraph()
        p_name.paragraph_format.space_after = Pt(2)
        r_name = p_name.add_run("ANTONIO GUTIÉRREZ JIMÉNEZ")
        r_name.font.name = 'Calibri'
        r_name.font.size = Pt(20)
        r_name.font.bold = True
        r_name.font.color.rgb = PRIMARY_COLOR

        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_after = Pt(6)
        r_title = p_title.add_run("Senior Business Development & Sales Leader | Payments, Acquiring & AI-Driven Infrastructure")
        r_title.font.name = 'Calibri'
        r_title.font.size = Pt(11)
        r_title.font.bold = True
        r_title.font.color.rgb = ACCENT_COLOR

        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(10)
        r_contact = p_contact.add_run("📍 Cancún, Q.R., Mexico   |   📞 +52 998 119 1903   |   ✉️ antoniogtzjimenez@gmail.com   |   🔗 linkedin.com/in/agjbusiness/")
        r_contact.font.name = 'Calibri'
        r_contact.font.size = Pt(9.5)
        r_contact.font.color.rgb = MUTED_COLOR

        # PROFESSIONAL SUMMARY
        add_heading("Professional Summary")
        p_sum = doc.add_paragraph()
        r_sum = p_sum.add_run(
            "Senior Business Development & B2B Sales Leader with 7+ years of experience accelerating commercial growth across Payments (PayTech), Merchant Acquiring, Financial Services, and High-Volume Retail in LATAM. Proven track record executing complex end-to-end enterprise sales cycles, negotiating API/ISV partnerships, and scaling TPV across middle-market and corporate accounts.\n\n"
            "Expert in outbound hunting, pipeline structuring, and financial ecosystem dynamics (Acquirers, Gateways, PSPs, Wallets, and A2A networks). Combines deep domain expertise in payment acceptance with advanced proficiency in Artificial Intelligence Tools, Autonomous AI Agents, and SalesTech automation to optimize pipeline velocity, account scoring, and commercial decision-making. Native Spanish, Professional English."
        )

        # CORE COMPETENCIES
        add_heading("Core Competencies & Domain Expertise")
        add_bullet("Merchant Acquiring, Payment Gateways, Payment Orchestration, A2A/Real-Time Payments (SPEI, PIX, PSE, Yape), Cross-Border Payments, Chargeback & Fraud Prevention.", bold_prefix="Payments & Financial Infrastructure: ")
        add_bullet("End-to-End Consultative Sales, Outbound Prospecting, C-Suite/VP Relationship Building, Multi-Stakeholder Contract Negotiation, Pipeline Structuring (Salesforce).", bold_prefix="Enterprise BD & Sales Execution: ")
        add_bullet("Technical-Commercial API integrations, ERP/POS/PMS Partnership Development, Solution Selling for Financial Institutions.", bold_prefix="API & ISV Ecosystem Partnerships: ")
        add_bullet("AI Agents & Prompt Engineering for RevOps/Sales Intelligence, Data Analytics (Power BI, Python), Automated Sales Workflows (Power Automate, CRM Orchestration).", bold_prefix="AI & Modern Sales Technology: ")

        # PROFESSIONAL EXPERIENCE
        add_heading("Professional Experience")

        # Co-Founder
        p_e1 = doc.add_paragraph()
        p_e1.paragraph_format.space_before = Pt(6)
        p_e1.paragraph_format.space_after = Pt(2)
        p_e1.paragraph_format.keep_with_next = True
        r_role = p_e1.add_run("Co-Founder  |  LATAM Payments & eCommerce Community")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e1.add_run("   •   05/2024 – Present   •   LATAM (Remote)")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Founded and lead an active community of 500+ payment professionals, merchant acquiring leaders, and e-commerce executives across Mexico and LATAM.")
        add_bullet("Leverage market intelligence and industry signals to identify emerging trends in digital payments, acquiring infrastructure, and cross-border orchestration.")

        # Fiserv
        p_e2 = doc.add_paragraph()
        p_e2.paragraph_format.space_before = Pt(8)
        p_e2.paragraph_format.space_after = Pt(2)
        p_e2.paragraph_format.keep_with_next = True
        r_role = p_e2.add_run("Business Advisor  |  Fiserv")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e2.add_run("   •   02/2025 – 10/2025   •   Cancún, Mexico")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Managed commercial development, retention, and churn mitigation for a portfolio of 80+ high-volume corporate merchants and enterprise accounts.")
        add_bullet("Built a Month-over-Month (MoM) transactional intelligence model using Power BI, enhancing portfolio health monitoring and identifying upsell/cross-sell signals.")
        add_bullet("Engineered an automated partner reactivation workflow leveraging Power Automate and Salesforce CRM, generating an average of +15 qualified commercial opportunities monthly through assigned banking alliances.")

        # Clip
        p_e3 = doc.add_paragraph()
        p_e3.paragraph_format.space_before = Pt(8)
        p_e3.paragraph_format.space_after = Pt(2)
        p_e3.paragraph_format.keep_with_next = True
        r_role = p_e3.add_run("Commercial Advisor – Middle Market / High Potential  |  Clip")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e3.add_run("   •   07/2021 – 02/2025   •   Cancún, Mexico")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Ranked #22 out of 184 Middle Market Executives nationwide (Top 12% national rank) in H1 2022. Consistently outperformed monthly volume quotas by over 280% ($2.8M to $5.8M MXN average monthly TPV vs. $1M quota), achieving #3 overall on the national performance podium.", bold_prefix="Top Performer Track Record: ")
        add_bullet("Executed a targeted commercial strategy on high-potential mid-market accounts, achieving an average TPV per deal of $555k MXN (60% higher than segment average), maximizing revenue while maintaining low onboarding friction.", bold_prefix="High-Value Portfolio Efficiency: ")
        add_bullet("Independently prospected, structured, and closed the highest-volume accounts in the regional portfolio, including a luxury yacht operator ($14.5M MXN YTD TPV) and an adventure tourism operator ($20.0M MXN YTD TPV).", bold_prefix="Enterprise Outbound Hunting: ")
        add_bullet("Led complex commercial negotiations and API/ISV payment gateway integrations with key ERP/POS platforms (Bistrosoft, Profitroom, Odoo), driving client retention with near-zero churn.", bold_prefix="API & Technical Integrations: ")
        add_bullet("Generated $69M MXN in cumulative TPV, with 75.3% of total volume self-generated through outbound hunting.", bold_prefix="Consolidated Portfolio Impact: ")

        # JTI
        p_e4 = doc.add_paragraph()
        p_e4.paragraph_format.space_before = Pt(8)
        p_e4.paragraph_format.space_after = Pt(2)
        p_e4.paragraph_format.keep_with_next = True
        r_role = p_e4.add_run("Account Executive · Southeast & Bajío  |  Japan Tobacco International (JTI)")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e4.add_run("   •   07/2018 – 12/2020   •   Cancún & Aguascalientes, Mexico")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Led regional territory expansion and Key Account Management (KAM) across HORECA and commercial channels, delivering a +40% YoY increase in Share of Opportunity in Cancún & Riviera Maya.")
        add_bullet("Expanded active client portfolio by +35%, securing direct commercial agreements with 100+ premium hotels and international hospitality brands.")
        add_bullet("Managed and directed a field commercial team of 3 sales representatives in the Bajío region.")

        # EDUCATION
        add_heading("Education")
        p_ed = doc.add_paragraph()
        p_ed.paragraph_format.space_after = Pt(2)
        r_ed = p_ed.add_run("Bachelor’s Degree in Commercial Relations (Licenciatura en Relaciones Comerciales)")
        r_ed.bold = True
        r_ed.font.color.rgb = PRIMARY_COLOR
        p_ed2 = doc.add_paragraph()
        p_ed2.paragraph_format.space_after = Pt(4)
        r_ed2 = p_ed2.add_run("Instituto Politécnico Nacional (IPN) — Mexico City   |   2014 – 2018 (Graduated / Titulado)")
        r_ed2.font.color.rgb = MUTED_COLOR
        r_ed2.font.size = Pt(9.5)

        # CERTIFICATIONS
        add_heading("Continuous Education & Certifications")
        add_bullet("McKinsey Forward Program (McKinsey.org · 120 hours) — Structured problem solving, adaptive leadership, and agile execution.")
        add_bullet("Growth 101 (Kurios · 30 hours) — Product-led growth frameworks and rapid experimentation.")
        add_bullet("Mastering Ventas (Sales Professional · 70 hours) — B2B sales playbooks, team structuring, and SalesTech stack automation.")
        add_bullet("Enterprise SDR & Prospecting (LATAM SDR Leaders · 16 hours) — Advanced outbound prospecting and enterprise account engagement.")

    else:
        # HEADER ES
        p_name = doc.add_paragraph()
        p_name.paragraph_format.space_after = Pt(2)
        r_name = p_name.add_run("ANTONIO GUTIÉRREZ JIMÉNEZ")
        r_name.font.name = 'Calibri'
        r_name.font.size = Pt(20)
        r_name.font.bold = True
        r_name.font.color.rgb = PRIMARY_COLOR

        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_after = Pt(6)
        r_title = p_title.add_run("Senior Business Development & Sales Leader | Pagos, Adquirencia e Infraestructura con IA")
        r_title.font.name = 'Calibri'
        r_title.font.size = Pt(11)
        r_title.font.bold = True
        r_title.font.color.rgb = ACCENT_COLOR

        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(10)
        r_contact = p_contact.add_run("📍 Cancún, Q.R., México   |   📞 +52 998 119 1903   |   ✉️ antoniogtzjimenez@gmail.com   |   🔗 linkedin.com/in/agjbusiness/")
        r_contact.font.name = 'Calibri'
        r_contact.font.size = Pt(9.5)
        r_contact.font.color.rgb = MUTED_COLOR

        # RESUMEN PROFESIONAL
        add_heading("Resumen Profesional")
        p_sum = doc.add_paragraph()
        r_sum = p_sum.add_run(
            "Líder Comercial Senior de Desarrollo de Negocios (B2B Sales) con más de 7 años de trayectoria acelerando el crecimiento de ingresos en Pagos Digitales (PayTech), Adquirencia, Servicios Financieros y Retail en LATAM. Con un récord demostrado ejecutando ciclos de venta complejos de punta a punta (End-to-End Enterprise), negociando alianzas tecnológicas vía API/ISV y escalando TPV en cuentas medianas y corporativas.\n\n"
            "Experto en prospección outbound, estructuración de pipeline en Salesforce y dinámicas del ecosistema financiero (Adquirentes, Pasarelas, PSPs, Wallets y redes A2A como SPEI, PIX, PSE, Yape). Combina un profundo conocimiento del mercado de pagos con un dominio avanzado de Herramientas de Inteligencia Artificial, Agentes Autónomos de IA y Automatización SalesTech para optimizar la velocidad del pipeline, el scoring de cuentas y la toma de decisiones comerciales. Español nativo, Inglés profesional."
        )

        # COMPETENCIAS CLAVE
        add_heading("Competencias Clave & Expertise de Dominio")
        add_bullet("Adquirencia (Merchant Acquiring), Pasarelas de Pago, Orquestación de Pagos, Pagos en Tiempo Real / A2A (SPEI, PIX, PSE, Yape), Pagos Transfronterizos (Cross-Border), Prevención de Fraude y Contracargos.", bold_prefix="Infraestructura de Pagos & Finanzas: ")
        add_bullet("Ventas Consultivas End-to-End, Prospección Outbound, Relaciones C-Level / VP, Negociación Compleja Multilado, Estructuración de Pipeline (Salesforce).", bold_prefix="Ejecución Comercial Enterprise: ")
        add_bullet("Integraciones Técnico-Comerciales vía API, Desarrollo de Alianzas con ERPs/POS/PMS, Soluciones para Instituciones Financieras.", bold_prefix="Alianzas Ecosistémicas API & ISV: ")
        add_bullet("Agentes de IA & Prompt Engineering para Inteligencia Comercial/RevOps, Analítica de Datos (Power BI, Python), Workflows Automatizados (Power Automate, CRM Orchestration).", bold_prefix="IA & Tecnología Comercial Moderna: ")

        # EXPERIENCIA PROFESIONAL
        add_heading("Experiencia Profesional")

        # Co-Founder
        p_e1 = doc.add_paragraph()
        p_e1.paragraph_format.space_before = Pt(6)
        p_e1.paragraph_format.space_after = Pt(2)
        p_e1.paragraph_format.keep_with_next = True
        r_role = p_e1.add_run("Co-Founder  |  Comunidad LATAM Payments & eCommerce")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e1.add_run("   •   05/2024 – Presente   •   LATAM (Remoto)")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Cofundé y coordino una comunidad activa de más de 500 profesionales de pagos digitales, adquirencia y comercio electrónico en México y LATAM.")
        add_bullet("Utilizo señales e inteligencia de mercado para identificar tendencias emergentes en pagos digitales, infraestructura de adquirencia y orquestación transfronteriza.")

        # Fiserv
        p_e2 = doc.add_paragraph()
        p_e2.paragraph_format.space_before = Pt(8)
        p_e2.paragraph_format.space_after = Pt(2)
        p_e2.paragraph_format.keep_with_next = True
        r_role = p_e2.add_run("Business Advisor  |  Fiserv")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e2.add_run("   •   02/2025 – 10/2025   •   Cancún, México")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Gestioné el desarrollo comercial, retención y mitigación de Churn para una cartera activa de más de 80 comercios corporativos (merchants) de gran escala.")
        add_bullet("Diseñé e implementé un modelo dinámico de monitoreo transaccional Month-over-Month (MoM) mediante Power BI, optimizando el seguimiento de salud de cartera e identificando señales de venta cruzada.")
        add_bullet("Desarrollé un workflow automatizado de reactivación mediante Power Automate e integraciones de Salesforce, generando un promedio de +15 oportunidades de negocio mensuales desde la alianza bancaria asignada.")

        # Clip
        p_e3 = doc.add_paragraph()
        p_e3.paragraph_format.space_before = Pt(8)
        p_e3.paragraph_format.space_after = Pt(2)
        p_e3.paragraph_format.keep_with_next = True
        r_role = p_e3.add_run("Asesor Comercial – Middle Market / High Potential  |  Clip")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e3.add_run("   •   07/2021 – 02/2025   •   Cancún, México")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Posicionado en el Top 12% nacional (Lugar #22 de 184 ejecutivos de Middle Market) en H1 2022. Superé las cuotas de volumen mensual asignadas por más del 280% de forma consistente ($2.8M a $5.8M MXN promedio frente a la meta de $1M), logrando el 3er lugar general del podio nacional.", bold_prefix="Récord de Top Performer: ")
        add_bullet("Ejecuté una estrategia comercial enfocada en cuentas medianas de alto potencial, logrando un TPV promedio por deal de $555k MXN (60% superior a la media del segmento), maximizando el volumen procesado con bajo costo operativo.", bold_prefix="Eficiencia de Cartera (High Value): ")
        add_bullet("Cerré de manera autónoma las cuentas de mayor volumen de la cartera en el sector turismo de lujo ($14.5M MXN YTD TPV en yates de lujo y $20.0M MXN YTD TPV en turismo de aventura).", bold_prefix="Cierre de Cuentas Enterprise (Outbound): ")
        add_bullet("Lideré negociaciones comerciales complejas e integraciones de pasarela de pagos vía API/ISV con sistemas clave (Bistrosoft, Profitroom, Odoo ERP), incrementando la retención de clientes con churn cercano a cero.", bold_prefix="Integraciones Tecnológicas & APIs: ")
        add_bullet("$69M MXN de TPV total acumulado, con un 75.3% del volumen total auto-generado vía Outbound.", bold_prefix="Resultado Acumulado de Cartera: ")

        # JTI
        p_e4 = doc.add_paragraph()
        p_e4.paragraph_format.space_before = Pt(8)
        p_e4.paragraph_format.space_after = Pt(2)
        p_e4.paragraph_format.keep_with_next = True
        r_role = p_e4.add_run("Account Executive · Southeast & Bajío  |  Japan Tobacco International (JTI)")
        r_role.bold = True
        r_role.font.color.rgb = PRIMARY_COLOR
        r_role.font.size = Pt(11)
        r_date = p_e4.add_run("   •   07/2018 – 12/2020   •   Cancún & Aguascalientes, México")
        r_date.font.color.rgb = MUTED_COLOR
        r_date.font.size = Pt(9.5)

        add_bullet("Lideré la expansión territorial y negociación de Key Accounts (KAM) en el sector HORECA, logrando un crecimiento del +40% YoY en Share of Opportunity en Cancún y Riviera Maya.")
        add_bullet("Expandí la cartera de clientes activos en un +35%, asegurando contratos de distribución directa con más de 100 hoteles premium y cadenas internacionales.")
        add_bullet("Coordiné y lideré un equipo comercial de 3 personas en campo para la región del Bajío.")

        # EDUCACIÓN
        add_heading("Educación")
        p_ed = doc.add_paragraph()
        p_ed.paragraph_format.space_after = Pt(2)
        r_ed = p_ed.add_run("Licenciatura en Relaciones Comerciales")
        r_ed.bold = True
        r_ed.font.color.rgb = PRIMARY_COLOR
        p_ed2 = doc.add_paragraph()
        p_ed2.paragraph_format.space_after = Pt(4)
        r_ed2 = p_ed2.add_run("Instituto Politécnico Nacional (IPN) — Ciudad de México   |   2014 – 2018 (Titulado)")
        r_ed2.font.color.rgb = MUTED_COLOR
        r_ed2.font.size = Pt(9.5)

        # CERTIFICACIONES
        add_heading("Educación Continua & Certificaciones")
        add_bullet("McKinsey Forward Program (McKinsey.org · 120 horas) — Resolución estructurada de problemas complejos, liderazgo adaptativo y ejecución ágil.")
        add_bullet("Growth 101 (Kurios · 30 horas) — Marcos de crecimiento acelerado y experimentación ágil de producto.")
        add_bullet("Mastering Ventas (Sales Professional · 70 horas) — Playbook de ventas B2B, estructura de equipos comerciales y automatización SalesTech.")
        add_bullet("Prospección Enterprise & SDR (LATAM SDR Leaders · 16 horas) — Estrategias avanzadas de prospección en frío y agendamiento de cuentas Enterprise.")

    doc.save(output_path)
    print(f"Generated styled Word CV: {output_path}")

out_dir = r"c:\Users\Antonio\OneDrive\Escritorio\CVs_2026"
create_styled_cv(os.path.join(out_dir, "CV_Antonio_Gutierrez_Yuno_Director.docx"), lang='en')
create_styled_cv(os.path.join(out_dir, "CV_Antonio_Gutierrez_Yuno_Director_ES.docx"), lang='es')
